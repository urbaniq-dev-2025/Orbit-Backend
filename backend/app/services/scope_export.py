"""
Scope Export Service

Handles exporting scopes to PDF, DOCX, and print formats.
"""

from __future__ import annotations

import json
import uuid
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Optional

from sqlalchemy.ext.asyncio import AsyncSession

from app.core.logging import get_logger
from app.models import Scope, ScopeSection
from app.services import scopes as scope_service

logger = get_logger(__name__)

# Export storage directory (in production, use S3 or similar)
EXPORT_STORAGE_DIR = Path("/tmp/orbit-exports")
EXPORT_STORAGE_DIR.mkdir(parents=True, exist_ok=True)


async def export_scope_to_pdf(
    session: AsyncSession,
    scope_id: uuid.UUID,
    user_id: uuid.UUID,
    *,
    include_sections: bool = True,
    template: str = "standard",
) -> tuple[bytes, str]:
    """
    Export scope to PDF format.
    Returns (pdf_bytes, filename)
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate
        from xml.sax.saxutils import escape
    except ImportError:
        logger.error("reportlab not installed. Install with: pip install reportlab")
        raise ValueError("PDF export requires reportlab library")

    scope = await scope_service.get_scope(session, scope_id, user_id, include_sections=include_sections)

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
    )
    story = []
    styles = getSampleStyleSheet()
    heading = styles["Heading1"]
    heading.spaceAfter = 12
    subheading = styles["Heading2"]
    body = styles["BodyText"]
    body.spaceAfter = 6

    # Title
    story.append(Paragraph(escape(scope.title or "Scope Document"), heading))
    if scope.description:
        story.append(Paragraph(escape(scope.description), body))

    # Scope metadata
    story.append(Paragraph(f"Status: {scope.status}", body))
    story.append(Paragraph(f"Progress: {scope.progress}%", body))
    if scope.due_date:
        story.append(Paragraph(f"Due Date: {scope.due_date.strftime('%Y-%m-%d')}", body))

    # Sections
    if include_sections and scope.sections:
        story.append(PageBreak())
        story.append(Paragraph("Sections", heading))
        for section in sorted(scope.sections, key=lambda x: x.order_index or 0):
            story.append(Paragraph(escape(section.title or "Untitled Section"), subheading))
            if section.content:
                # Split content into paragraphs
                content_paragraphs = section.content.split("\n\n")
                for para in content_paragraphs:
                    if para.strip():
                        story.append(Paragraph(escape(para.strip()), body))

    # Scope document JSON (if available)
    if scope.scope_document_json:
        try:
            scope_doc = json.loads(scope.scope_document_json)
            story.append(PageBreak())
            story.append(Paragraph("Scope Details", heading))

            # Executive Summary
            if "executive_summary" in scope_doc:
                es = scope_doc["executive_summary"]
                story.append(Paragraph("Executive Summary", subheading))
                if es.get("overview"):
                    story.append(Paragraph(escape(es["overview"]), body))
                if es.get("key_points"):
                    story.append(Paragraph("Key Points:", body))
                    for point in es["key_points"]:
                        story.append(Paragraph(f"• {escape(str(point))}", body))

            # Modules
            if "modules" in scope_doc and scope_doc["modules"]:
                story.append(PageBreak())
                story.append(Paragraph("Modules", subheading))
                for module in scope_doc["modules"]:
                    story.append(Paragraph(escape(module.get("name", "Unnamed Module")), subheading))
                    if module.get("description"):
                        story.append(Paragraph(escape(module["description"]), body))

            # Features
            if "features" in scope_doc and scope_doc["features"]:
                story.append(PageBreak())
                story.append(Paragraph("Features", subheading))
                for feature in scope_doc["features"]:
                    story.append(Paragraph(escape(feature.get("name", "Unnamed Feature")), subheading))
                    if feature.get("summary"):
                        story.append(Paragraph(escape(feature["summary"]), body))
                    if feature.get("priority"):
                        story.append(Paragraph(f"Priority: {escape(str(feature['priority']))}", body))

        except json.JSONDecodeError:
            logger.warning(f"Failed to parse scope_document_json for scope {scope_id}")

    doc.build(story)
    buffer.seek(0)
    pdf_bytes = buffer.read()

    filename = f"scope_{scope_id}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.pdf"
    return pdf_bytes, filename


async def export_scope_to_docx(
    session: AsyncSession,
    scope_id: uuid.UUID,
    user_id: uuid.UUID,
    *,
    include_sections: bool = True,
    template: str = "standard",
) -> tuple[bytes, str]:
    """
    Export scope to DOCX format.
    Returns (docx_bytes, filename)
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise ValueError("DOCX export requires python-docx library")

    scope = await scope_service.get_scope(session, scope_id, user_id, include_sections=include_sections)

    doc = Document()
    
    # Title
    title = doc.add_heading(scope.title or "Scope Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if scope.description:
        doc.add_paragraph(scope.description)

    # Scope metadata
    doc.add_paragraph(f"Status: {scope.status}")
    doc.add_paragraph(f"Progress: {scope.progress}%")
    if scope.due_date:
        doc.add_paragraph(f"Due Date: {scope.due_date.strftime('%Y-%m-%d')}")

    # Sections
    if include_sections and scope.sections:
        doc.add_page_break()
        doc.add_heading("Sections", 1)
        for section in sorted(scope.sections, key=lambda x: x.order_index or 0):
            doc.add_heading(section.title or "Untitled Section", 2)
            if section.content:
                # Split content into paragraphs
                content_paragraphs = section.content.split("\n\n")
                for para in content_paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())

    # Scope document JSON (if available)
    if scope.scope_document_json:
        try:
            scope_doc = json.loads(scope.scope_document_json)
            doc.add_page_break()
            doc.add_heading("Scope Details", 1)

            # Executive Summary
            if "executive_summary" in scope_doc:
                es = scope_doc["executive_summary"]
                doc.add_heading("Executive Summary", 2)
                if es.get("overview"):
                    doc.add_paragraph(es["overview"])
                if es.get("key_points"):
                    doc.add_paragraph("Key Points:")
                    for point in es["key_points"]:
                        doc.add_paragraph(f"• {point}", style="List Bullet")

            # Modules
            if "modules" in scope_doc and scope_doc["modules"]:
                doc.add_page_break()
                doc.add_heading("Modules", 2)
                for module in scope_doc["modules"]:
                    doc.add_heading(module.get("name", "Unnamed Module"), 3)
                    if module.get("description"):
                        doc.add_paragraph(module["description"])

            # Features
            if "features" in scope_doc and scope_doc["features"]:
                doc.add_page_break()
                doc.add_heading("Features", 2)
                for feature in scope_doc["features"]:
                    doc.add_heading(feature.get("name", "Unnamed Feature"), 3)
                    if feature.get("summary"):
                        doc.add_paragraph(feature["summary"])
                    if feature.get("priority"):
                        doc.add_paragraph(f"Priority: {feature['priority']}")

        except json.JSONDecodeError:
            logger.warning(f"Failed to parse scope_document_json for scope {scope_id}")

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.read()

    filename = f"scope_{scope_id}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.docx"
    return docx_bytes, filename


async def get_scope_for_print(
    session: AsyncSession,
    scope_id: uuid.UUID,
    user_id: uuid.UUID,
    *,
    include_sections: bool = True,
) -> dict:
    """
    Get scope data optimized for printing.
    Returns a structured dictionary ready for print rendering.
    """
    scope = await scope_service.get_scope(session, scope_id, user_id, include_sections=include_sections)

    # Build print-optimized structure
    print_data = {
        "title": scope.title,
        "description": scope.description,
        "status": scope.status,
        "progress": scope.progress,
        "dueDate": scope.due_date.isoformat() if scope.due_date else None,
        "createdAt": scope.created_at.isoformat() if scope.created_at else None,
        "updatedAt": scope.updated_at.isoformat() if scope.updated_at else None,
        "sections": [],
    }

    # Add sections
    if include_sections and scope.sections:
        for section in sorted(scope.sections, key=lambda x: x.order_index or 0):
            print_data["sections"].append({
                "title": section.title,
                "content": section.content,
                "sectionType": section.section_type,
                "orderIndex": section.order_index or 0,
            })

    # Add scope document JSON if available
    if scope.scope_document_json:
        try:
            scope_doc = json.loads(scope.scope_document_json)
            print_data["scopeDocument"] = scope_doc
        except json.JSONDecodeError:
            logger.warning(f"Failed to parse scope_document_json for scope {scope_id}")

    return print_data


async def save_export_file(scope_id: uuid.UUID, filename: str, file_bytes: bytes) -> str:
    """
    Save export file to storage and return the file path.
    In production, this would upload to S3 and return a signed URL.
    """
    file_path = EXPORT_STORAGE_DIR / f"{scope_id}_{filename}"
    file_path.write_bytes(file_bytes)
    
    # Return relative path for download endpoint
    return f"/api/scopes/{scope_id}/exports/{filename}"
